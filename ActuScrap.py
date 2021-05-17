#!/usr/bin/env python
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
import time
import os
from openpyxl import Workbook
from selenium.webdriver.common.by import By
from docx import Document

option = webdriver.ChromeOptions()
option.add_argument("--incognito")
option.add_argument("--start-maximised")

browser = webdriver.Chrome("./chromedriver", options=option)
Article = input("Bienvenue : Veuillez coller l'URL de l'article du parisien que vous souhaitez récupérer")
browser.get(Article)
time.sleep(2)
browser.maximize_window()
bouton = browser.find_element_by_id('didomi-notice-agree-button')
bouton.click()
Texte = browser.find_element_by_xpath("//div[@class='article-section margin_bottom_article']").text
Titre = browser.find_element_by_xpath("//h1[@class='title_xl col margin_bottom_headline']").text




document = Document()
document.add_heading(Titre,0)
document.add_paragraph(Texte)
document.save('article.docx')


